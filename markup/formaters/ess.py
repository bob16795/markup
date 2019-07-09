"""
handles .ess files according to
#olympus.sandhills.edu/english/wordguide/mlaformat.html
"""
from docx.enum import text
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def setup(doc):
    """
    formats a .ess file
    """
    doc.sections[0].right_margin = Inches(1)
    doc.sections[0].left_margin = Inches(1)
    doc.sections[0].top_margin = Inches(1)
    doc.sections[0].bottom_margin = Inches(1)

    norm = doc.styles['Normal']
    norm.paragraph_format.line_spacing = 2
    norm.font.name = 'Times New Roman'
    norm.font.size = Pt(12)

    head = doc.styles.add_style('head', 1)
    head.paragraph_format.line_spacing = 2
    head.font.name = 'Times New Roman'
    head.font.size = Pt(12)

    til = doc.styles.add_style("til", 1)
    til.paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.CENTER
    til.paragraph_format.line_spacing = 2
    til.font.size = Pt(12)

def header(doc, title, tags):
    """
    Adds a heading to a .doc file
    """
    print(f" + Ess {title}")
    paragraph = doc.sections[0].header.add_paragraph(f"\t{tags['Lastname']} ")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'PAGE \\*Arabic \\* MERGEFORMAT'   # change 1-3 depending on heading levels you need
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    doc.add_paragraph(tags['Author'], "head")

    doc.add_paragraph(tags['Prof'], "head")

    doc.add_paragraph(tags['Class'], "head")

    doc.add_paragraph(tags['Date'], "head")

    doc.add_paragraph(title, "til")
