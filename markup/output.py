from docx.shared import RGBColor
from docx import Document
from markup.decorators import Formater
from pygments.formatter import Formatter
import subprocess
import tempfile
import os
import re


@Formater
class terminal():
    def outer_init(self, out, prop):
        self.out = out

    def outer_add(self, out):
        self.out += out
        return self

    def fmt_init(self, **options):
        Formatter.__init__(self, **options)

    def fmt_format(self, tokensource, outfile):
        lastval = ''
        lasttype = None
        for ttype, value in tokensource:
            if ttype == lasttype:
                lastval += value
            else:
                if lastval:
                    outfile += lastval
                lastval = value
                lasttype = ttype
        if lastval:
            outfile += lastval.replace("\n", "")

    @staticmethod
    def add_new_line():
        return "\n"

    @staticmethod
    def add_text(text):
        return text

    @staticmethod
    def emph_text(text):
        return text

    @staticmethod
    def start():
        return ""

    @staticmethod
    def bold_text(text):
        return text

    @staticmethod
    def add_header_1(text):
        text = "\n" + text + "\n" + (len(text) * "=") + "\n"
        return text

    @staticmethod
    def add_header_2(text):
        text = "\n" + text + "\n" + (len(text) * "-") + "\n"
        return text

    @staticmethod
    def add_header_3(text):
        text = "\n" + text + "\n"
        return text

    @staticmethod
    def start_list():
        return "\n-"

    @staticmethod
    def start_code():
        return "\n"

    @staticmethod
    def code_line(text):
        return ">  " + text + "\n"

    @staticmethod
    def end_code():
        return ""

    @staticmethod
    def start_list_1(l):
        return "\n-"

    @staticmethod
    def start_list_2(l):
        return "\n  -"

    @staticmethod
    def start_list_3(l):
        return "\n    -"

    @staticmethod
    def end_list(l):
        return "\n"

    @staticmethod
    def emph_list_text(text):
        return text

    @staticmethod
    def bold_list_text(text):
        return text

    @staticmethod
    def add_list_text(text):
        return text

    @staticmethod
    def end(out):
        out += "\n"
        return out.out.encode()

    @staticmethod
    def tag(text):
        return f"link: [{text}]"


@Formater
class html():
    @staticmethod
    def outer_init(self, out, prop):
        self.out = out

    @staticmethod
    def outer_add(self, out):
        self.out += out
        return self

    @staticmethod
    def fmt_init(self, **options):
        Formatter.__init__(self, **options)

        self.styles = {}

        for token, style in self.style:
            start = end = ''

            if style['color']:
                start += '<font color="#%s">' % style['color']
                end = '</font>' + end
            if style['bold']:
                start += '<b>'
                end = '</b>' + end
            if style['italic']:
                start += '<i>'
                end = '</i>' + end
            if style['underline']:
                start += '<u>'
                end = '</u>' + end
            self.styles[token] = (start, end)

    @staticmethod
    def fmt_format(self, tokensource, outfile):
        lastval = ''
        lasttype = None
        outfile += '<pre>'
        for ttype, value in tokensource:
            while ttype not in self.styles:
                ttype = ttype.parent
            if ttype == lasttype:
                lastval += value
            else:
                if lastval:
                    stylebegin, styleend = self.styles[lasttype]
                    outfile += stylebegin + lastval + styleend
                lastval = value
                lasttype = ttype
        if lastval:
            stylebegin, styleend = self.styles[lasttype]
            outfile += stylebegin + lastval + styleend
            outfile += '</pre>\n'

    @staticmethod
    def start():
        return "<!DOCTYPE html>\n<html>\n<head>\n<link rel=\"stylesheet\" href=\"main.css\">\n</head>\n<body>\n<div id=\"page\">"

    @staticmethod
    def add_text(text):
        return text

    @staticmethod
    def add_new_line():
        return "<br/>"

    @staticmethod
    def emph_text(text):
        return "<em>" + text + "</em>"

    @staticmethod
    def bold_text(text):
        return "<b>" + text + "</b>"

    @staticmethod
    def add_header_1(text):
        return "<h1>" + text + "</h1>"

    @staticmethod
    def add_header_2(text):
        return "<h2>" + text + "</h2>"

    @staticmethod
    def add_header_3(text):
        return "<h3>" + text + "</h3>"

    @staticmethod
    def start_code():
        return "\n"

    @staticmethod
    def code_line(text):
        return "" + text + "\n"

    @staticmethod
    def end_code():
        return ""

    @staticmethod
    def start_list():
        return "<ul><li>"

    @staticmethod
    def start_list_1(l):
        if l == 1:
            return "</li><li>"
        if l == 2:
            return "</li></ul><li>"
        if l == 3:
            return "</li></ul></ul><li>"

    @staticmethod
    def start_list_2(l):
        if l == 1:
            return "</li><ul><li>"
        if l == 2:
            return "</li><li>"
        if l == 3:
            return "</li></ul><li>"

    @staticmethod
    def start_list_3(l):
        if l == 1:
            return "</li><ul><ul><li>"
        if l == 2:
            return "</li><ul><li>"
        if l == 3:
            return "</li><li>"

    @staticmethod
    def end_list(l):
        if l == 1:
            return "</li></ul>"
        if l == 2:
            return "</li></ul></ul>"
        if l == 3:
            return "</li></ul></ul></ul>"

    @staticmethod
    def emph_list_text(text):
        return "<em>" + text + "</em>"

    @staticmethod
    def bold_list_text(text):
        return "<b>" + text + "</b>"

    @staticmethod
    def add_list_text(text):
        return text

    @staticmethod
    def end(out):
        out += "</div>\n</body>\n</html>"
        return out.out.encode()

    @staticmethod
    def tag(text):
        if "](" in text:
            text = text.split("](")
            link = text[0]
            text = text[-1]
        else:
            text = text.split(": ")
            link = text[0]
            text = text[-1]
        return f"<a href={text}> {link} </a><br/>"


@Formater
class docx():
    @staticmethod
    def outer_init(self, out, prop):
        if "template" in prop:
            self.doc = Document(prop["template"])
        else:
            self.doc = Document()

    @staticmethod
    def outer_add(self, out):
        for i in out:
            ttype = i.split(": ")[0]
            text = i.split(": ")[-1]
            if ttype == "\n":
                self.p = None
            elif ttype == "H1":
                self.doc.add_heading(text, 1)
                self.p = None
            elif ttype == "H2":
                self.doc.add_heading(text, 2)
                self.p = None
            elif ttype == "H3":
                self.doc.add_heading(text, 3)
                self.p = None
            elif ttype == "BOLD":
                if self.p == None:
                    self.p = self.doc.add_paragraph()
                self.r = self.p.add_run(text)
                self.r.bold = True
            elif ttype == "EMPH":
                if self.p == None:
                    self.p = self.doc.add_paragraph()
                self.r = self.p.add_run(text)
                self.r.italic = True
            else:
                if self.p == None:
                    self.p = self.doc.add_paragraph()
                self.r = self.p.add_run(text)
        return self

    @staticmethod
    def fmt_init(self, **options):
        Formatter.__init__(self, **options)

        # create a dict of (start, end) tuples that wrap the
        # value of a token so that we can use it in the format
        # method later
        self.styles = {}

        # we iterate over the `_styles` attribute of a style item
        # that contains the parsed style values.
        for token, style in self.style:
            # a style item is a tuple in the following form:
            # colors are readily specified in hex: 'RRGGBB'
            if style['color'] == None:
                color = RGBColor.from_string("000000")
            else:
                color = RGBColor.from_string(style['color'])
            self.styles[token] = (color, style['bold'],
                                  style['italic'], style['underline'])

    @staticmethod
    def fmt_format(self, tokensource, outfile):
        # lastval is a string we use for caching
        # because it's possible that an lexer yields a number
        # of consecutive tokens with the same token type.
        # to minimize the size of the generated html markup we
        # try to join the values of same-type tokens here
        lastval = ''
        lasttype = None
        doc = outfile
        self.p = doc.add_paragraph(style="Source Code")
        # wrap the whole output with <pre>

        for ttype, value in tokensource:
            # if the token type doesn't exist in the stylemap
            # we try it with the parent of the token type
            # eg: parent of Token.Literal.String.Double is
            # Token.Literal.String
            while ttype not in self.styles:
                ttype = ttype.parent
            if ttype == lasttype:
                # the current token type is the same of the last
                # iteration. cache it
                lastval += value
            else:
                # not the same token as last iteration, but we
                # have some data in the buffer. wrap it with the
                # defined style and write it to the output file
                if lastval:
                    color, bold, italic, underline = self.styles[lasttype]
                    r = self.p.add_run(lastval)
                    r.font.color.rgb = color
                    r.font.bold = bold
                    r.font.italic = italic
                    r.font.underline = underline
                # set lastval/lasttype to current5 values
                lastval = value
                lasttype = ttype

        # if something is left in the buffer, write it to the
        if lastval:
            color, bold, italic, underline = self.styles[lasttype]
            r = self.p.add_run(lastval.replace("\n", ""))
            r.font.color.rgb = color
            r.font.bold = bold
            r.font.italic = italic
            r.font.underline = underline

    @staticmethod
    def start():
        return []

    @staticmethod
    def add_text(text):
        return [text]

    @staticmethod
    def add_new_line():
        return ["\n"]

    @staticmethod
    def emph_text(text):
        return [f"EMPH: {text}"]

    @staticmethod
    def bold_text(text):
        return [f"BOLD: {text}"]

    @staticmethod
    def add_header_1(text):
        return [f"H1: {text}"]

    @staticmethod
    def add_header_2(text):
        return [f"H2: {text}"]

    @staticmethod
    def add_header_3(text):
        return [f"H3: {text}"]

    @staticmethod
    def start_code():
        return ["\n"]

    @staticmethod
    def code_line(text):
        return [text, "\n"]

    @staticmethod
    def end_code():
        return []

    @staticmethod
    def start_list():
        return []

    @staticmethod
    def start_list_1(l):
        return [f"L1: "]

    @staticmethod
    def start_list_2(l):
        return [f"L2: "]

    @staticmethod
    def start_list_3(l):
        return [f"L3: "]

    @staticmethod
    def end_list(l):
        return []

    @staticmethod
    def emph_list_text(text):
        return [f"EMPH: {text}"]

    @staticmethod
    def bold_list_text(text):
        return [f"BOLD: {text}"]

    @staticmethod
    def add_list_text(text):
        return [f"{text}"]

    @staticmethod
    def end(out):
        out.doc.save("/tmp/tmp.docx")
        with open("/tmp/tmp.docx", "rb") as f:
            text = f.read()
        return text

    @staticmethod
    def tag(text):
        if "](" in text:
            text = text.split("](")
            link = text[0]
            text = text[-1]
        else:
            text = text.split(": ")
            link = text[0]
            text = text[-1]
        return [f"<a href={text}> {link} </a><br/>"]


@Formater
class pdf_groff():
    """
    args:
        author: the author of the document
        title: the title of the document
        title_page: if defined will add a title page to the document
        title_head: the level of the first heading

        TODO: margins
    """
    @staticmethod
    def outer_init(self, out, prop):
        self.pp = False
        self.title_heading_level = 0
        self.author = ""
        self.title = ""
        if "author" in prop:
            self.author = prop["author"]
        if "title" in prop:
            self.title = prop["title"]
            if "title_page" in prop:
                out += "\n()TTL()\n.bp\n"
        if "title_head" in prop:
            self.title_heading_level = int(prop["title_head"])

        out = out.replace("()TTL()", self.title)
        out = out.replace("()AUT()", self.author)
        self.out = out

    @staticmethod
    def outer_add(self, out):
        out = out.replace("()HL1()", str(self.title_heading_level))
        out = out.replace("()HL2()", str(self.title_heading_level + 1))
        out = out.replace("()HL3()", str(self.title_heading_level + 2))
        if not out == "":
            if out[0] == "\n" or out[:3] == ".SH" or \
                    out[:3] == "LI;":
                if out[:3] == "LI;":
                    out = out[3:]
                elif self.out[-1] != "\n":
                    self.out += "\n"
                self.pp = False
            elif self.pp == False:
                if self.out[-1] != "\n":
                    self.out += "\n"
                self.pp = True
                self.out += ".PP\n"
            if not out == "\n":
                self.out += out
        return self

    @staticmethod
    def fmt_init(self, **options):
        Formatter.__init__(self, **options)
        self.styles = {}
        for token, style in self.style:
            if style['color'] == None:
                color = "000000"
            else:
                color = style['color']
            self.styles[token] = (color, style['bold'],
                                  style['italic'], style['underline'])

    @staticmethod
    def fmt_format(self, tokensource, outfile):
        lastval = ''
        lasttype = None
        outfile += "\n.in 0\n.fcolor gray\n.LP\n"
        indent = 0
        for ttype, value in tokensource:
            if not(value.strip(" ")):
                indent = len(value)-len(value.strip(" "))
            if lastval.strip(" "):
                color = self.styles[lasttype][0]
                val = lastval.replace('\n', f'\n.in {float(indent/6)}c\n')
                outfile += f".defcolor pyg rgb #{color}\n.gcolor pyg\n{val}\n.gcolor\n"
            # set lastval/lasttype to current values
            lastval = value
            lasttype = ttype
        if lastval.strip(" "):
            color = self.styles[lasttype][0]
            val = lastval.replace('\n', f'\n.in {float(indent/6)}c\n')
            outfile += f".defcolor pyg rgb #{color}\n.gcolor pyg\n{val}\n.gcolor\n"
        # set lastval/lasttype to current values
        lastval = value
        lasttype = ttype
        outfile += f".fcolor\n"

    @staticmethod
    def add_new_line():
        return f"\n"

    @staticmethod
    def add_text(text):
        if text.strip(" ") != "":
            return f"{text}\n"
        return ""

    @staticmethod
    def emph_text(text):
        return f".UL \"{text}\"\n"

    @staticmethod
    def start():
        # \\X'papersize=5.5i,8.5i'\n
        # \n.nr PO .3i\n.nr LL 6.4i\n.nr FM .5i\n.nr HM .3i\n.nr LT 7.4i
        return "\n.OH '''%'\n.EH '''%'\n.color 1\n.OF '()AUT()'''\n.EF '''()TTL()'\n"

    @staticmethod
    def bold_text(text):
        return f".B {text}\n"

    @staticmethod
    def add_header_1(text):
        return f".OH '%'-{text}-''\n.EH ''-{text}-'%'\n.bp\n.NH ()HL1()\n{text}\n.XS\n.B\n{text}\n.XE\n.PP\n"

    @staticmethod
    def add_header_2(text):
        return f".NH ()HL2()\n{text}\n.XS\n\t{text}\n.XE\n.PP\n"

    @staticmethod
    def add_header_3(text):
        return f".NH ()HL3()\n{text}\n.XS\n\t\t{text}\n.XE\n.PP\n"

    @staticmethod
    def start_list():
        return ""

    @staticmethod
    def start_code():
        return "LI;\n.LP\n"

    @staticmethod
    def code_line(text):
        return text.replace("\\TS", "#\\.RS\n")\
            .replace("\\TE", "#\\.RE\n") + "\n"

    @staticmethod
    def end_code():
        return ""

    @staticmethod
    def start_list_1(l):
        if l == 1:
            return ".IP \\(bu 2\n"
        if l == 2:
            return ".RE\n.IP \\(bu 2\n"
        if l == 3:
            return ".RE\n.RE\n.IP \\(bu 2\n"

    @staticmethod
    def start_list_2(l):
        if l == 1:
            return ".RS\n.IP \\(bu 2\n"
        if l == 2:
            return ".IP \\(bu 2\n"
        if l == 3:
            return ".RE\n.IP \\(bu 2\n"

    @staticmethod
    def start_list_3(l):
        if l == 1:
            return ".RS\n.RS\n.IP \\(bu 2\n"
        if l == 2:
            return ".RS\n.IP \\(bu 2\n"
        if l == 3:
            return ".IP \\(bu 2\n"

    @staticmethod
    def end_list(l):
        if l == 1:
            return ""
        if l == 2:
            return ".RE\n"
        if l == 3:
            return ".RE\n.RE\n"

    @staticmethod
    def emph_list_text(text):
        return f"LI;{text}"

    @staticmethod
    def bold_list_text(text):
        return f"LI;{text}"

    @staticmethod
    def add_list_text(text):
        return f"LI;{text}"

    @staticmethod
    def end(out):
        out += ".OH '%'-Table Of Contents-''\n.EH ''-Table Of Contents-'%'\n.de TOC\n.MC 200p .3i\n.SH\nTable Of Contents\n..\n.TC"
        out.out = out.out.replace("\n\n", "\n")
        o = subprocess.Popen(f"groff -Tpdf -dpaper=a4 -P-pa4 -ms".split(" "),
                             stdin=subprocess.PIPE, stdout=subprocess.PIPE)
        return o.communicate(input=out.out.encode())[0]

    @staticmethod
    def tag(text):
        if "](" in text:
            text = text.split("](")
            link = text[0]
            text = text[-1]
        else:
            text = text.split(": ")
            link = text[0]
            text = text[-1]
        if link == "COL":
            # TODO: calculate coloums automatically
            return f"\n.MC {text}i\n"
        if link == "CPT":
            if text[-1] == "!":
                return f"\n.bp\n.NH 0\n{text[:-1]}\n"
            else:
                return f".OH '%'-Table Of Contents-''\n.EH ''-Table Of Contents-'%'\n.de TOC\n.MC 200p .3i\n.SH\nTable Of Contents\n..\n.TC\n.bp\n.NH 0\n{text}\n.rm toc*div\n.rm toc*num\n"
        return f"{text}\n"

# TODO: chapter numbers
@Formater
class pdf_latex():
    @staticmethod
    def outer_init(self, out, prop):
        self.pp = False
        self.title_heading_level = 0
        self.author = ""
        self.title = ""
        """
        \def\l@subsubsubsection{\@dottedtocline{4}{7em}{4em}}
        \def\l@paragraph{\@dottedtocline{5}{10em}{5em}}
        \def\l@subparagraph{\@dottedtocline{6}{14em}{6em}}
        """
        out += "\def\l@subsubsubsection{\@dottedtocline{()HL1()}{7em}{4em}}\n\def\l@paragraph{\@dottedtocline{()HL2()}{10em}{5em}}\n\def\l@subparagraph{\@dottedtocline{()HL3()}{14em}{6em}}\n\\begin{document}"
        if "author" in prop:
            self.author = prop["author"]
        if "title" in prop:
            self.title = prop["title"]
            if "title_page" in prop:
                """
                \\begin{titlepage}
                    \\begingroup
                    \\vspace*{0.12\\textheight}
                    \\hspace*{0.3\\textwidth}
                    \\hspace*{0.3\\textwidth}
                    {\\Huge ()TTL()}\\par
                    \\vspace*{0.36\\textheight}
                    {\\large ()AUT()}
                    \\vfill
                    \\endgroup
                \\end{titlepage}
                """
                out += "\n\\begin{titlepage}\n\\begingroup\n\\vspace*{0.12\\textheight}\n\\hspace*{0.3\\textwidth}\n\\hspace*{0.3\\textwidth}\n{\\Huge ()TTL()}\\par\n\\vspace*{0.36\\textheight}\n{\\large ()AUT()}\n\\vfill\n\\endgroup\n\\end{titlepage}\n"
        if "title_head" in prop:
            self.title_heading_level = int(prop["title_head"])
        if "author" in prop:
            self.author = prop["author"]
        if "paper_size" in prop:
            self.paper_size = prop["paper_size"]
        out = out.replace("()HL1()", str(self.title_heading_level + 0))
        out = out.replace("()HL2()", str(self.title_heading_level + 1))
        out = out.replace("()HL3()", str(self.title_heading_level + 2))
        out = out.replace("()TTL()", self.title)
        if self.author:
            out = out.replace("()AUT()", "by: " + self.author)
        else:
            out = out.replace("()AUT()", "")
        if self.paper_size:
            out = out.replace("()PPR()", self.paper_size)
        else:
            out = out.replace("()PPR()", "a4paper")
        out += "\\begin{multicols}{1}"
        self.out = out

    @staticmethod
    def outer_add(self, out):
        if self.author:
            out = out.replace("()AUT()", "by: " + self.author)
        else:
            out = out.replace("()AUT()", "")
        if self.paper_size:
            out = out.replace("()PPR()", self.paper_size)
        else:
            out = out.replace("()PPR()", "a4paper")
        if self.title:
            out = out.replace("()TTL()", self.title)
        else:
            out = out.replace("()TTL()", "")
        if self.out[-1] != "\n":
            self.out += "\n"
        if not out == "\n":
            self.out += out
        return self

    @staticmethod
    def fmt_init(self, **options):
        Formatter.__init__(self, **options)
        self.styles = {}
        for token, style in self.style:
            if style['color'] == None:
                color = "000000"
            else:
                color = style['color']
            self.styles[token] = (color, style['bold'],
                                  style['italic'], style['underline'])

    @staticmethod
    def fmt_format(self, tokensource, outfile):
        lastval = ''
        lasttype = None
        indent = 0
        outfile += "\\begin{flushleft}\n{\\texttt{\n"
        for ttype, value in tokensource:
            if not(value.strip(" ")):
                indent = len(value)-len(value.strip(" "))
            if lastval.strip(" "):
                color_r = int(self.styles[lasttype][0], 16) & int("FF0000", 16)
                color_g = int(self.styles[lasttype][0], 16) & int("00FF00", 16)
                color_b = int(self.styles[lasttype][0], 16) & int("0000FF", 16)
                color_r = str(float(color_r + 1)/float(256))
                color_g = str(float(color_g + 1)/float(256))
                color_b = str(float(color_b + 1)/float(256))
                color = "{" + f"{color_r}, {color_g}, {color_b}" + "}"
                ind = "{" + str(indent*8) + "pt}"
                val = lastval.replace("\\", "{\\textbackslash}").replace('{', '\\{').replace('$', '\\$').replace('}', '\\}').replace("\n",f"\\\\\n\\setlength\\parindent{ind}\n") #.replace('[', '\\[').replace(']', '\\]')
                code = "{code}"
                rgb = "{rgb}"
                outfile += f"\\definecolor{code}{rgb}{color}" + \
                "\n{" + f"\\color{code} {val}" + "}"
            lastval = value
            lasttype = ttype
        if lastval.strip(" "):
            color_r = int(self.styles[lasttype][0], 16) & int("FF0000", 16)
            color_g = int(self.styles[lasttype][0], 16) & int("FF00", 16)
            color_b = int(self.styles[lasttype][0], 16) & int("FF", 16)
            color_r = str(float(color_r + 1)/float(256))
            color_g = str(float(color_g + 1)/float(256))
            color_b = str(float(color_b + 1)/float(256))
            color = "{" + f"{color_r}, {color_g}, {color_b}" + "}"
            val = lastval.replace("\\", "{\\textbackslash}").replace('{', '\\{').replace('$', '\\$').replace('}', '\\}').replace("\n","\\\\\n") #.replace('[', '\\[').replace(']', '\\]')
            code = "{code}"
            rgb = "{rgb}"
            outfile += f"\\definecolor{code}{rgb}{color}" + \
                "\n{" + f"\\color{code} {val}" + "}"
        outfile += "}}\n\\end{flushleft}\n"
    @staticmethod
    def add_new_line():
        return f"\n"

    @staticmethod
    def add_text(text):
        if text.strip(" ") != "":
            return f"{text}\n"
        return ""

    @staticmethod
    def emph_text(text):
        text = "{"+text+"}"
        return f"\\emph{text}\n"

    @staticmethod
    def start():
        return "\\documentclass{book}\n\\usepackage{multicol}\n\\usepackage[()PPR()]{geometry}\n\\usepackage{xcolor}\n"

    @staticmethod
    def bold_text(text):
        text = "{"+text+"}"
        return f"\\textbf{text}\n"

    @staticmethod
    def add_header_1(text):
        text = "{"+text+"}"
        return f"\\section{text}\n"

    @staticmethod
    def add_header_2(text):
        text = "{"+text+"}"
        return f"\\subsection{text}\n"

    @staticmethod
    def add_header_3(text):
        text = "{"+text+"}"
        return f"\\subsubsection{text}\n"

    @staticmethod
    def start_list():
        return "\\begin{itemize}"

    @staticmethod
    def start_code():
        return ""

    @staticmethod
    def code_line(text):
        return text+"\n"

    @staticmethod
    def end_code():
        return ""

    @staticmethod
    def start_list_1(l):
        if l == 1:
            return ""
        if l == 2:
            return "\\end{itemize}\n"
        if l == 3:
            return "\\end{itemize}\n\\end{itemize}\n"

    @staticmethod
    def start_list_2(l):
        if l == 1:
            return "\\begin{itemize}\n"
        if l == 2:
            return ""
        if l == 3:
            return "\\end{itemize}\n"

    @staticmethod
    def start_list_3(l):
        if l == 1:
            return "\\begin{itemize}\n\\begin{itemize}\n"
        if l == 2:
            return "\\begin{itemize}\n"
        if l == 3:
            return ""

    @staticmethod
    def end_list(l):
        if l == 1:
            return "\\end{itemize}\n"
        if l == 2:
            return "\\end{itemize}\n\\end{itemize}\n"
        if l == 3:
            return "\\end{itemize}\n\\end{itemize}\n\\end{itemize}\n"

    @staticmethod
    def emph_list_text(text):
        return f"\\item {text}\n"

    @staticmethod
    def bold_list_text(text):
        return f"\\item {text}\n"

    @staticmethod
    def add_list_text(text):
        return f"\\item {text}\n"

    # TODO: fix dependency on user being me
    @staticmethod
    def end(out):
        tmpdir = "/tmp/"
        if os.name == "nt":
            tmpdir = "C:\\Users\\Preston.precourt\\Downloads\\"
        out += "\\end{multicols}\n\\end{document}"
        out.out = out.out.replace("&", "\\&").replace("#", "\\#").replace(
            "\\n", "{\\textbackslash}n").replace("_", "\\_").replace("|", "\\|")
        tempin = tempfile.NamedTemporaryFile(dir=f"{tmpdir}", delete=False)
        tempin.write(out.out.encode())
        tempin_name = tmpdir + \
            tempfile.gettempprefix() + tempin.name.split("tmp")[-1]
        tempin.close()
        path = tmpdir
        try:
            if os.name == "nt":
                o = subprocess.Popen(f"C:\\Users\\Preston.precourt\\AppData\\Local\\Programs\\texlive\\texlive\\2019\\bin\\win32\\pdflatex.exe -output-directory {path} {tempin_name}".split(" "),
                                 stdin=subprocess.PIPE, stdout=subprocess.PIPE)
            else:
                o = subprocess.Popen(f"pdflatex -output-directory {path} {tempin_name}".split(" "),
                                 stdin=subprocess.PIPE, stdout=subprocess.PIPE)
            out = o.communicate()
            tempout = open(f"{tempin_name}.pdf", 'r+b')
            pdf = tempout.read()
            tempout.close()
        except:
            print("error")
            if out:
                print(out)
            pdf = ""
        for file in sorted(os.listdir(tmpdir)):
            if re.search("^tmp.", file):
                os.remove(tmpdir + file)
        return pdf

    @staticmethod
    def tag(text):
        if "](" in text:
            text = text.split("](")
            link = text[0].strip(" ")
            text = text[-1].strip(" ")
        else:
            text = text.split(":")
            link = text[0].strip(" ")
            text = text[-1].strip(" ")
        if link == "COL":
            return "\\end{multicols}\\begin{multicols}{" + text + "}"
        if link == "CPT":
            text = "{"+text+"}"
            text = "[\n\\chapter" + text + "\n]\n"
            return text
        return f"{text}\n"

# TODO: finish code
# TODO: chapter numbers
@Formater
class latex():
    @staticmethod
    def outer_init(self, out, prop):
        self.pp = False
        self.title_heading_level = 0
        self.author = ""
        self.title = ""
        """
        \def\l@subsubsubsection{\@dottedtocline{4}{7em}{4em}}
        \def\l@paragraph{\@dottedtocline{5}{10em}{5em}}
        \def\l@subparagraph{\@dottedtocline{6}{14em}{6em}}
        """
        out += "\def\l@subsubsubsection{\@dottedtocline{()HL1()}{7em}{4em}}\n\def\l@paragraph{\@dottedtocline{()HL2()}{10em}{5em}}\n\def\l@subparagraph{\@dottedtocline{()HL3()}{14em}{6em}}\n\\begin{document}"
        if "author" in prop:
            self.author = prop["author"]
        if "title" in prop:
            self.title = prop["title"]
            if "title_page" in prop:
                """
                \\begin{titlepage}
                    \\begingroup
                    \\vspace*{0.12\\textheight}
                    \\hspace*{0.3\\textwidth}
                    \\hspace*{0.3\\textwidth}
                    {\\Huge ()TTL()}\\par
                    \\vspace*{0.36\\textheight}
                    {\\large ()AUT()}
                    \\vfill
                    \\endgroup
                \\end{titlepage}
                """
                out += "\n\\begin{titlepage}\n\\begingroup\n\\vspace*{0.12\\textheight}\n\\hspace*{0.3\\textwidth}\n\\hspace*{0.3\\textwidth}\n{\\Huge ()TTL()}\\par\n\\vspace*{0.36\\textheight}\n{\\large ()AUT()}\n\\vfill\n\\endgroup\n\\end{titlepage}\n"
        if "table_of_contents" in prop:
            out += "\n\\tableofcontents\n\\newpage\n"
        if "title_head" in prop:
            self.title_heading_level = int(prop["title_head"])
        if "author" in prop:
            self.author = prop["author"]
        if "paper_size" in prop:
            self.paper_size = prop["paper_size"]
        out = out.replace("()HL1()", str(self.title_heading_level + 0))
        out = out.replace("()HL2()", str(self.title_heading_level + 1))
        out = out.replace("()HL3()", str(self.title_heading_level + 2))
        out = out.replace("()TTL()", self.title)
        if self.author:
            out = out.replace("()AUT()", "by: " + self.author)
        else:
            out = out.replace("()AUT()", "")
        if self.paper_size:
            out = out.replace("()PPR()", self.paper_size)
        else:
            out = out.replace("()PPR()", "a4paper")
        out += "\\begin{multicols}{1}"
        self.out = out

    @staticmethod
    def outer_add(self, out):
        if self.author:
            out = out.replace("()AUT()", "by: " + self.author)
        else:
            out = out.replace("()AUT()", "")
        if self.paper_size:
            out = out.replace("()PPR()", self.paper_size)
        else:
            out = out.replace("()PPR()", "a4paper")
        if self.title:
            out = out.replace("()TTL()", self.title)
        else:
            out = out.replace("()TTL()", "")
        if self.out[-1] != "\n":
            self.out += "\n"
        if not out == "\n":
            self.out += out
        return self

    @staticmethod
    def fmt_init(self, **options):
        Formatter.__init__(self, **options)
        self.styles = {}
        for token, style in self.style:
            if style['color'] == None:
                color = "000000"
            else:
                color = style['color']
            self.styles[token] = (color, style['bold'],
                                  style['italic'], style['underline'])

    @staticmethod
    def fmt_format(self, tokensource, outfile):
        lastval = ''
        lasttype = None
        indent = 0
        outfile += "\\begin{flushleft}\n{\\texttt{\n"
        for ttype, value in tokensource:
            if not(value.strip(" ")):
                indent = len(value)-len(value.strip(" "))
            if lastval.strip(" "):
                color_r = int(self.styles[lasttype][0], 16) & int("FF0000", 16)
                color_g = int(self.styles[lasttype][0], 16) & int("00FF00", 16)
                color_b = int(self.styles[lasttype][0], 16) & int("0000FF", 16)
                color_r = str(float(color_r + 1)/float(256))
                color_g = str(float(color_g + 1)/float(256))
                color_b = str(float(color_b + 1)/float(256))
                color = "{" + f"{color_r}, {color_g}, {color_b}" + "}"
                ind = "{" + str(indent*8) + "pt}"
                val = lastval.replace("\\", "{\\textbackslash}").replace('{', '\\{').replace('$', '\\$').replace('}', '\\}').replace("\n",f"\\\\\n\\setlength\\parindent{ind}\n") #.replace('[', '\\[').replace(']', '\\]')
                code = "{code}"
                rgb = "{rgb}"
                outfile += f"\\definecolor{code}{rgb}{color}" + \
                "\n{" + f"\\color{code} {val}" + "}"
            lastval = value
            lasttype = ttype
        if lastval.strip(" "):
            color_r = int(self.styles[lasttype][0], 16) & int("FF0000", 16)
            color_g = int(self.styles[lasttype][0], 16) & int("FF00", 16)
            color_b = int(self.styles[lasttype][0], 16) & int("FF", 16)
            color_r = str(float(color_r + 1)/float(256))
            color_g = str(float(color_g + 1)/float(256))
            color_b = str(float(color_b + 1)/float(256))
            color = "{" + f"{color_r}, {color_g}, {color_b}" + "}"
            val = lastval.replace("\\", "{\\textbackslash}").replace('{', '\\{').replace('$', '\\$').replace('}', '\\}').replace("\n","\\\\\n") #.replace('[', '\\[').replace(']', '\\]')
            code = "{code}"
            rgb = "{rgb}"
            outfile += f"\\definecolor{code}{rgb}{color}" + \
                "\n{" + f"\\color{code} {val}" + "}"
        outfile += "}}\n\\end{flushleft}\n"
    @staticmethod
    def add_new_line():
        return f"\n"

    @staticmethod
    def add_text(text):
        if text.strip(" ") != "":
            return f"{text}\n"
        return ""

    @staticmethod
    def emph_text(text):
        text = "{"+text+"}"
        return f"\\emph{text}\n"

    @staticmethod
    def start():
        return "\\documentclass{book}\n\\usepackage{multicol}\n\\usepackage[()PPR()]{geometry}\n\\usepackage{xcolor}\n"

    @staticmethod
    def bold_text(text):
        text = "{"+text+"}"
        return f"\\textbf{text}\n"

    @staticmethod
    def add_header_1(text):
        text = "{"+text+"}"
        return f"\\section{text}\n"

    @staticmethod
    def add_header_2(text):
        text = "{"+text+"}"
        return f"\\subsection{text}\n"

    @staticmethod
    def add_header_3(text):
        text = "{"+text+"}"
        return f"\\subsubsection{text}\n"

    @staticmethod
    def start_list():
        return "\\begin{itemize}"

    @staticmethod
    def start_code():
        return ""

    @staticmethod
    def code_line(text):
        return text+"\n"

    @staticmethod
    def end_code():
        return ""

    @staticmethod
    def start_list_1(l):
        if l == 1:
            return ""
        if l == 2:
            return "\\end{itemize}\n"
        if l == 3:
            return "\\end{itemize}\n\\end{itemize}\n"

    @staticmethod
    def start_list_2(l):
        if l == 1:
            return "\\begin{itemize}\n"
        if l == 2:
            return ""
        if l == 3:
            return "\\end{itemize}\n"

    @staticmethod
    def start_list_3(l):
        if l == 1:
            return "\\begin{itemize}\n\\begin{itemize}\n"
        if l == 2:
            return "\\begin{itemize}\n"
        if l == 3:
            return ""

    @staticmethod
    def end_list(l):
        if l == 1:
            return "\\end{itemize}\n"
        if l == 2:
            return "\\end{itemize}\n\\end{itemize}\n"
        if l == 3:
            return "\\end{itemize}\n\\end{itemize}\n\\end{itemize}\n"

    @staticmethod
    def emph_list_text(text):
        return f"\\item {text}\n"

    @staticmethod
    def bold_list_text(text):
        return f"\\item {text}\n"

    @staticmethod
    def add_list_text(text):
        return f"\\item {text}\n"

    # TODO: fix dependency on user being me
    @staticmethod
    def end(out):
        out += "\\end{multicols}\n\\end{document}"
        out.out = out.out.replace("&", "\\&").replace("#", "\\#").replace(
            "\\n", "{\\textbackslash}n").replace("_", "\\_").replace("|", "\\|")
        return out.out.encode("utf-8")

    @staticmethod
    def tag(text):
        if "](" in text:
            text = text.split("](")
            link = text[0].strip(" ")
            text = text[-1].strip(" ")
        else:
            text = text.split(":")
            link = text[0].strip(" ")
            text = text[-1].strip(" ")
        if link == "COL":
            return "\\end{multicols}\\begin{multicols}{" + text + "}"
        if link == "CPT":
            text = "{"+text+"}"
            text = "[\n\\chapter" + text + "\n]\n"
            return text
        return f"{text}\n"
